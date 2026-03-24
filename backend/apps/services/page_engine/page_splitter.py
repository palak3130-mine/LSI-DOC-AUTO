import fitz
from docx import Document
import os
import uuid


def split_pdf_by_page(file_path, output_dir):
    """
    Split PDF into individual pages
    """
    pages = []

    doc = fitz.open(file_path)

    for page_number in range(len(doc)):
        new_doc = fitz.open()
        new_doc.insert_pdf(doc, from_page=page_number, to_page=page_number)

        filename = f"page_{page_number+1}_{uuid.uuid4().hex}.pdf"
        output_path = os.path.join(output_dir, filename)

        new_doc.save(output_path)
        new_doc.close()

        pages.append({
            "page_number": page_number + 1,
            "path": output_path
        })

    return pages


def split_docx_by_page(file_path, output_dir):
    """
    Split DOCX by paragraphs (pseudo pages)
    """
    doc = Document(file_path)

    pages = []
    page = []
    page_number = 1

    # Approx paragraphs per page
    MAX_PARAGRAPHS = 20

    for para in doc.paragraphs:
        page.append(para.text)

        if len(page) >= MAX_PARAGRAPHS:
            filename = f"page_{page_number}_{uuid.uuid4().hex}.docx"
            output_path = os.path.join(output_dir, filename)

            new_doc = Document()

            for line in page:
                new_doc.add_paragraph(line)

            new_doc.save(output_path)

            pages.append({
                "page_number": page_number,
                "path": output_path
            })

            page = []
            page_number += 1

    # Last page
    if page:
        filename = f"page_{page_number}_{uuid.uuid4().hex}.docx"
        output_path = os.path.join(output_dir, filename)

        new_doc = Document()

        for line in page:
            new_doc.add_paragraph(line)

        new_doc.save(output_path)

        pages.append({
            "page_number": page_number,
            "path": output_path
        })

    return pages


def split_document_by_page(file_path, output_dir):
    """
    Auto detect and split document
    """
    if file_path.endswith(".pdf"):
        return split_pdf_by_page(file_path, output_dir)

    elif file_path.endswith(".docx"):
        return split_docx_by_page(file_path, output_dir)

    else:
        raise ValueError("Unsupported file format")