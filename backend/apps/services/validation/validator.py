import os
from docx import Document


def validate_pages(updated_pages):
    """
    Validate updated pages
    """

    print("🔍 Validating Updated Pages")

    valid_pages = []

    for page in updated_pages:

        page_path = page["path"]

        if not os.path.exists(page_path):
            print(f"❌ Missing Page: {page_path}")
            continue

        try:
            doc = Document(page_path)

            text = "\n".join([p.text for p in doc.paragraphs])

            if not text.strip():
                print(f"⚠️ Empty Page: {page_path}")
                continue

            valid_pages.append(page)

        except Exception as e:
            print(f"❌ Invalid Page {page_path}: {e}")

    print(f"✅ {len(valid_pages)} pages validated")

    return valid_pages


def validate_final_document(output_path):
    """
    Validate final stitched document
    """

    print("🔍 Validating Final Document")

    if not os.path.exists(output_path):
        raise Exception("Final document missing")

    try:
        doc = Document(output_path)

        text = "\n".join([p.text for p in doc.paragraphs])

        if not text.strip():
            raise Exception("Final document is empty")

    except Exception as e:
        raise Exception(f"Final document invalid: {e}")

    print("✅ Final document validated")

    return True