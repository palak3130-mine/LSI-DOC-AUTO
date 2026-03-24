import fitz  # PyMuPDF
from docx import Document as DocxDocument


def parse_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""

    for page in doc:
        text += page.get_text()

    return text


def parse_docx(file_path):
    doc = DocxDocument(file_path)
    text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)

    return "\n".join(text)


def parse_document(file_path):
    if file_path.endswith(".pdf"):
        return parse_pdf(file_path)

    elif file_path.endswith(".docx"):
        return parse_docx(file_path)

    else:
        raise ValueError("Unsupported file format")