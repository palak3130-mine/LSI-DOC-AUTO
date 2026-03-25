from docx import Document


def replace_text_preserve_format(paragraph, find_text, replace_text):
    """
    Safe text replacement without breaking formatting
    """

    if find_text not in paragraph.text:
        return False

    inline = paragraph.runs

    for run in inline:
        if find_text in run.text:
            run.text = run.text.replace(find_text, replace_text)
            return True

    # fallback (multi-run match)
    full_text = paragraph.text

    if find_text in full_text:
        new_text = full_text.replace(find_text, replace_text)

        # Only update first run
        inline[0].text = new_text

        # Clear others
        for run in inline[1:]:
            run.text = ""

        return True

    return False


def replace_text_in_table(table, find_text, replace_text):

    replaced = False

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_text_preserve_format(
                    paragraph,
                    find_text,
                    replace_text
                ):
                    replaced = True

    return replaced


def apply_changes(doc_path, changes, output_path):

    print("🚀 Starting Merge Engine...")

    doc = Document(doc_path)

    total_updates = 0

    for change in changes:

        find_text = change.get("find")
        replace_text = change.get("replace")

        if not find_text or not replace_text:
            continue

        replaced = False

        # paragraphs
        for paragraph in doc.paragraphs:
            if replace_text_preserve_format(
                paragraph,
                find_text,
                replace_text
            ):
                replaced = True

        # tables
        for table in doc.tables:
            if replace_text_in_table(
                table,
                find_text,
                replace_text
            ):
                replaced = True

        if replaced:
            total_updates += 1
            print(f"✅ Updated: {find_text[:50]}...")
        else:
            print(f"❌ Not found: {find_text[:50]}...")

    print(f"🎯 Total Updates: {total_updates}")

    doc.save(output_path)

    print("✅ Merge Engine Completed")

    return output_path